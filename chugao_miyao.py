import streamlit as st
import json
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.schema import HumanMessage, SystemMessage
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.callbacks.streamlit import StreamlitCallbackHandler
from langchain.chains import SequentialChain, LLMChain
import os
from typing import Dict, Any, List
import logging
import sys
from docx import Document
import io
import base64
from PyPDF2 import PdfReader
from PIL import Image
import fitz  # PyMuPDF
# 配置日志记录
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
from queue import Queue
from threading import Thread
import time
from queue import Empty
logger = logging.getLogger(__name__)
from langchain.callbacks.base import BaseCallbackHandler
from miyao import KeyManager
from datetime import datetime, timedelta

# 记录程序启动
logger.info("程序开始运行")

# 只在第一次运行时替换 sqlite3
if 'sqlite_setup_done' not in st.session_state:
    try:
        logger.info("尝试设置 SQLite")
        __import__('pysqlite3')
        sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
        st.session_state.sqlite_setup_done = True
        logger.info("SQLite 设置成功")
    except Exception as e:
        logger.error(f"SQLite 设置错误: {str(e)}")
        st.session_state.sqlite_setup_done = True


class PromptTemplates:
    def __init__(self):
        # 定义示例数据作为字符串
        self.default_templates = {
            'transcript_role': """
            # 角色
            你是专业的成绩单分析师，擅长从成绩单中提取关键信息并进行分析。
            """,
            
            'transcript_task': """
            分析学生的成绩单，提取以下信息：
            1. 学生的GPA和成绩分布情况
            2. 主要课程的成绩表现
            3. 学术优势和劣势
            4. 成绩趋势（是否有进步或下滑）
            5. 与申请专业相关课程的表现
            """,
            
            'transcript_output': """
            成绩单分析:
                GPA和总体表现: [GPA和总体成绩分布]
                主要课程成绩: [列出主要课程及成绩]
                学术优势: [分析学生的学术优势]
                学术劣势: [分析学生的学术劣势]
                成绩趋势: [分析成绩的变化趋势]
                与申请专业相关性: [分析与申请专业相关课程的表现]
            """,
            
            'consultant_role1': """
            # 角色
            你是资深留学顾问，精通学生背景分析和各国院校招生政策。
            """,
            
            'output_format1': """
            学生背景分析: 
                核心亮点: 亮点1，亮点2，亮点3...
                需要加强的方面: 需要加强的方面1，需要加强的方面2...
            申请策略: 
                国家与专业分析: 对目标国家和专业招生偏好的简要分析
                推荐写作方向: 方向1，方向2...
                核心卖点: 如何突出学生的优势并与专业匹配
            """,
            
            'consultant_task1': """
            根据选校方案先判断是否已选校，如果已选校，则结合选校方案进行后续分析
            分析学生的个人陈述表，提取关键信息与亮点
            如果有成绩单分析，结合成绩单分析结果进行综合评估
            根据申请国家和专业确定PS的写作大方向
            评估学生背景与目标专业的匹配度
            制定个性化文书策略，确定核心卖点
            """,
            
            'consultant_role2': """
            # 角色
            你是结构化思维与创意写作专家，擅长内容规划和素材创作。
            """,
            
            'output_format2': """
            文书框架: 
                整体结构: 文书整体结构概述
                段落规划: 
                    段落目的: 这段要达成的目标
                    核心内容: 应包含的关键信息
                    素材建议: 
                    需要补充的内容: 具体需要补充什么类型的素材
                    补充例子: 具体的素材示例
                    与专业关联: 如何将此素材与申请专业关联
                其他段落: 其他段落规划
            """,
            
            'consultant_task2': """
            根据选校方案先判断是否已选校，如果已选校，则结合选校方案进行后续分析
            设计PS的整体框架和段落结构
            为每个段落规划内容要点和与专业的关联
            直接提供具体素材补充建议和实例
            确保补充素材与学生背景一致且符合申请专业需求
            """
        }
        
        # 初始化 session_state 中的模板
        if 'templates' not in st.session_state:
            st.session_state.templates = self.default_templates.copy()

    def get_template(self, template_name: str) -> str:
        return st.session_state.templates.get(template_name, "")

    def update_template(self, template_name: str, new_content: str) -> None:
        st.session_state.templates[template_name] = new_content

    def reset_to_default(self):
        st.session_state.templates = self.default_templates.copy()

class TranscriptAnalyzer:
    def __init__(self, api_key: str, prompt_templates: PromptTemplates):
        # 使用OpenRouter API访问模型
        self.llm = ChatOpenAI(
            temperature=0.7,
            model=st.secrets["TRANSCRIPT_MODEL"],  # 从secrets中获取模型名称
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1",
            streaming=True
        )
        self.prompt_templates = prompt_templates
    
    def extract_images_from_pdf(self, pdf_bytes):
        """从PDF中提取图像"""
        try:
            images = []
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                # 将页面直接转换为图像
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_bytes = pix.tobytes("png")
                # 将图像编码为base64字符串
                img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                images.append(img_base64)
            
            return images
        except Exception as e:
            logger.error(f"提取PDF图像时出错: {str(e)}")
            return []
    
    def analyze_transcript(self, pdf_bytes, school_plan: str) -> Dict[str, Any]:
        try:
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
                
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行分析
            def run_analysis():
                try:
                    # 提取PDF中的图像
                    images = self.extract_images_from_pdf(pdf_bytes)
                    
                    if not images:
                        message_queue.put("无法从PDF中提取图像，请检查文件格式。")
                        return
                    
                    # 构建提示词
                    system_prompt = f"{self.prompt_templates.get_template('transcript_role')}\n\n" \
                                   f"任务:\n{self.prompt_templates.get_template('transcript_task')}\n\n" \
                                   f"请按照以下格式输出:\n{self.prompt_templates.get_template('transcript_output')}"
                    
                    # 将图像转换为文本描述
                    image_descriptions = [f"[图像 {i+1}: 成绩单页面]" for i in range(len(images))]
                    image_text = "\n".join(image_descriptions)
                    
                    user_prompt = f"选校方案：\n{school_plan}\n\n" \
                                 f"请分析以下成绩单图像，提取关键信息并进行分析。\n\n" \
                                 f"成绩单包含以下页面：\n{image_text}"
                    
                    # 创建提示模板
                    prompt = ChatPromptTemplate.from_messages([
                        ("system", system_prompt),
                        ("human", user_prompt)
                    ])
                    
                    # 调用LLM进行分析
                    chain = LLMChain(llm=self.llm, prompt=prompt)
                    result = chain.run(
                        {},
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    
                    message_queue.put("\n\n成绩单分析完成！")
                    thread.result = result
                    return result
                    
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"成绩单分析错误: {str(e)}")
                    thread.exception = e
                    raise e
            
            # 启动线程
            thread = Thread(target=run_analysis)
            thread.start()
            
            # 使用 st.write_stream 显示流式输出
            output_container = st.empty()
            with output_container:
                full_response = st.write_stream(token_generator())
            
            # 等待线程完成
            thread.join()
            
            # 获取结果
            if hasattr(thread, "exception") and thread.exception:
                raise thread.exception
            
            logger.info("成绩单分析完成")
            
            return {
                "status": "success",
                "transcript_analysis": full_response
            }
                
        except Exception as e:
            logger.error(f"成绩单分析错误: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }

class BrainstormingAgent:
    def __init__(self, api_key: str, prompt_templates: PromptTemplates):
        self.llm = ChatOpenAI(
            temperature=0.7,
            model=st.secrets["OPENROUTER_MODEL"],
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1",
            streaming=True
        )
        self.prompt_templates = prompt_templates
        self.setup_chains()

    def setup_chains(self):
        # Profile Strategist Chain
        strategist_prompt = ChatPromptTemplate.from_messages([
            ("system", f"{self.prompt_templates.get_template('consultant_role1')}\n\n"
                      f"任务:\n{self.prompt_templates.get_template('consultant_task1')}\n\n"
                      f"请按照以下格式输出:\n{self.prompt_templates.get_template('output_format1')}"),
            ("human", "选校方案：\n{school_plan}\n\n"
                     "请分析以下学生个人陈述：\n\n"
                     "个人陈述：\n{document_content}")
        ])
        
        self.strategist_chain = LLMChain(
            llm=self.llm,
            prompt=strategist_prompt,
            output_key="strategist_analysis",
            verbose=True
        )

        # Content Creator Chain - 更新提示词以包含成绩单分析和自定义需求
        creator_prompt = ChatPromptTemplate.from_messages([
            ("system", f"{self.prompt_templates.get_template('consultant_role2')}\n\n"
                      f"任务:\n{self.prompt_templates.get_template('consultant_task2')}\n\n"
                      f"请按照以下格式输出:\n{self.prompt_templates.get_template('output_format2')}"),
            ("human", "选校方案：\n{school_plan}\n\n"
                     "基于第一阶段的分析结果：\n{strategist_analysis}\n\n"
                     "成绩单分析结果：\n{transcript_analysis}\n\n"
                     "额外定制需求：\n{custom_requirements}\n\n"
                     "请创建详细的内容规划。")
        ])
        
        self.creator_chain = LLMChain(
            llm=self.llm,
            prompt=creator_prompt,
            output_key="creator_output",
            verbose=True
        )

    def process_strategist(self, document_content: str, school_plan: str, transcript_analysis: str = "") -> Dict[str, Any]:
        try:
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器，继承自 BaseCallbackHandler
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
                
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行LLM
            def run_llm():
                try:
                    result = self.strategist_chain(
                        {
                            "document_content": document_content, 
                            "school_plan": school_plan,
                            "transcript_analysis": transcript_analysis
                        },
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    # 将结果存储在线程对象中
                    thread.result = result
                    message_queue.put("\n\n分析完成！")
                    return result
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"Strategist processing error: {str(e)}")
                    thread.exception = e
                    raise e
            
            # 启动线程
            thread = Thread(target=run_llm)
            thread.start()
            
            # 使用 st.write_stream 显示流式输出
            output_container = st.empty()
            with output_container:
                full_response = st.write_stream(token_generator())
            
            # 等待线程完成
            thread.join()
            
            # 获取结果
            if hasattr(thread, "exception") and thread.exception:
                raise thread.exception
            
            logger.info("Strategist analysis completed successfully")
            
            # 从 full_response 中提取分析结果
            return {
                "status": "success",
                "strategist_analysis": full_response
            }
                
        except Exception as e:
            logger.error(f"Strategist processing error: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }
    def process_creator(self, strategist_analysis: str, school_plan: str, transcript_analysis: str = "", custom_requirements: str = "无定制需求") -> Dict[str, Any]:
        try:
            # 创建一个队列用于流式输出
            message_queue = Queue()
            
            # 创建自定义回调处理器，继承自 BaseCallbackHandler
            class QueueCallbackHandler(BaseCallbackHandler):
                def __init__(self, queue):
                    self.queue = queue
                    super().__init__()
            
                def on_llm_new_token(self, token: str, **kwargs) -> None:
                    self.queue.put(token)
            
            # 创建一个生成器函数，用于流式输出
            def token_generator():
                while True:
                    try:
                        token = message_queue.get(block=False)
                        yield token
                    except Empty:
                        if not thread.is_alive() and message_queue.empty():
                            break
                    time.sleep(0.01)
            
            # 在单独的线程中运行LLM
            def run_llm():
                try:
                    result = self.creator_chain(
                        {
                            "strategist_analysis": strategist_analysis,
                            "school_plan": school_plan,
                            "transcript_analysis": transcript_analysis,
                            "custom_requirements": custom_requirements
                        },
                        callbacks=[QueueCallbackHandler(message_queue)]
                    )
                    # 将结果存储在队列中
                    message_queue.put("\n\n规划完成！")
                    return result
                except Exception as e:
                    message_queue.put(f"\n\n错误: {str(e)}")
                    logger.error(f"Creator processing error: {str(e)}")
                    raise e
            
            # 启动线程
            thread = Thread(target=run_llm)
            thread.start()
            
            # 使用 st.write_stream 显示流式输出
            output_container = st.empty()
            with output_container:
                full_response = st.write_stream(token_generator())
            
            # 等待线程完成
            thread.join()
            
            # 获取结果
            if hasattr(thread, "_exception") and thread._exception:
                raise thread._exception
            
            logger.info("Creator analysis completed successfully")
            
            return {
                "status": "success",
                "creator_output": full_response
            }
                
        except Exception as e:
            logger.error(f"Creator processing error: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }


def add_custom_css():
    st.markdown("""
    <style>
    /* 标题样式 */
    h1, h2, h3 {
        color: #1e3a8a;
        font-weight: 600;
    }
    
    /* 卡片样式 */
    .stTabs [data-baseweb="tab-panel"] {
        background-color: white;
        border-radius: 10px;
        padding: 20px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.12), 0 1px 2px rgba(0,0,0,0.24);
        margin-top: 10px;
    }
    
    /* 按钮样式 */
    .stButton>button {
        background-color: #1e3a8a;
        color: white;
        border-radius: 5px;
        padding: 10px 20px;
        font-weight: 500;
        border: none;
        width: 100%;
    }
    
    .stButton>button:hover {
        background-color: #2e4a9a;
    }
    
    /* 输入框样式 */
    .stTextInput>div>div>input, .stTextArea>div>div>textarea {
        border-radius: 5px;
        border: 1px solid #ddd;
    }
    
    /* 文件上传区域样式 */
    .stFileUploader>div>button {
        background-color: #f1f3f9;
        color: #1e3a8a;
        border: 1px dashed #1e3a8a;
        border-radius: 5px;
    }
    
    /* 成功消息样式 */
    .stSuccess {
        background-color: #d1fae5;
        color: #065f46;
        padding: 10px;
        border-radius: 5px;
    }
    
    /* 警告消息样式 */
    .stWarning {
        background-color: #fef3c7;
        color: #92400e;
        padding: 10px;
        border-radius: 5px;
    }
    
    /* 错误消息样式 */
    .stError {
        background-color: #fee2e2;
        color: #b91c1c;
        padding: 10px;
        border-radius: 5px;
    }
    
    /* 下拉选择框样式 */
    .stSelectbox>div>div {
        border-radius: 5px;
        border: 1px solid #ddd;
    }
    
    /* 页面标题样式 */
    .page-title {
        text-align: center;
        font-size: 2rem;
        margin-bottom: 20px;
        color: #1e3a8a;
        font-weight: bold;
    }
    
    /* 卡片容器样式 */
    .card-container {
        background-color: white;
        border-radius: 10px;
        padding: 20px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.12), 0 1px 2px rgba(0,0,0,0.24);
        margin-bottom: 20px;
        width: 100%;
    }
    
    /* 分隔线样式 */
    hr {
        margin-top: 20px;
        margin-bottom: 20px;
        border: 0;
        border-top: 1px solid #eee;
    }
    
    /* 模型信息样式 */
    .model-info {
        background-color: #f0f7ff;
        padding: 8px 12px;
        border-radius: 5px;
        margin-top: 10px;
        margin-bottom: 15px;
        display: inline-block;
        font-size: 0.9rem;
    }
    
    /* 表格样式优化 */
    .dataframe {
        width: 100%;
        border-collapse: collapse;
    }
    
    .dataframe th {
        background-color: #f1f3f9;
        padding: 8px;
    }
    
    .dataframe td {
        padding: 8px;
        border-bottom: 1px solid #eee;
    }
    
    
    
    /* 调整列宽度 */
    .column-adjust {
        padding: 0 5px !important;
    }
    
    /* 强制展开器内容宽度 */
    .streamlit-expanderContent {
        width: 100% !important;
    }
    </style>
    """, unsafe_allow_html=True)


def read_docx(file_bytes):
    """读取 Word 文档内容，包括表格，并去除重复内容"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        content_set = set()  # 用于存储已处理的内容，避免重复
        full_text = []
        
        # 读取普通段落
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and text not in content_set:  # 只添加非空且未重复的内容
                content_set.add(text)
                full_text.append(text)
        
        # 读取表格内容
        for table in doc.tables:
            table_content = []
            header_row = []
            
            # 获取表头（第一行）
            if table.rows:
                for cell in table.rows[0].cells:
                    header_text = cell.text.strip()
                    if header_text:
                        header_row.append(header_text)
            
            # 处理表格内容（从第二行开始）
            for row_idx, row in enumerate(table.rows[1:], 1):
                row_content = {}
                for col_idx, cell in enumerate(row.cells):
                    if col_idx < len(header_row):  # 确保有对应的表头
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content[header_row[col_idx]] = cell_text
                
                if row_content:  # 只添加非空行
                    formatted_row = " | ".join([f"{header}: {value}" 
                                              for header, value in row_content.items()])
                    if formatted_row not in content_set:  # 避免重复内容
                        content_set.add(formatted_row)
                        table_content.append(formatted_row)
            
            if table_content:
                full_text.extend(table_content)
        
        # 使用换行符连接所有文本
        result = "\n".join(full_text)
        logger.info(f"成功读取文档内容，包含 {len(doc.tables)} 个表格")
        return result
    except Exception as e:
        logger.error(f"读取 Word 文档时出错: {str(e)}")
        return None

# 添加新的验证相关函数
def init_session_state():
    """初始化会话状态变量"""
    if 'auth_attempts' not in st.session_state:
        st.session_state.auth_attempts = 0
    if 'is_authenticated' not in st.session_state:
        st.session_state.is_authenticated = False
    if 'last_attempt_time' not in st.session_state:
        st.session_state.last_attempt_time = None
    if 'blocked_until' not in st.session_state:
        st.session_state.blocked_until = None

def get_client_info():
    """获取客户端信息"""
    try:
        # 获取 Streamlit 提供的客户端信息
        client_info = {
            'query_string': st.experimental_get_query_params(),
            'browser_info': st.experimental_user.browser,
            'time': datetime.now().isoformat()
        }
        return client_info
    except:
        return {}

def verify_access():
    """验证访问权限"""
    init_session_state()
    
    # 如果已经通过验证，直接返回True
    if st.session_state.is_authenticated:
        return True
    
    # 检查是否在封禁期
    if st.session_state.blocked_until:
        if datetime.now() < st.session_state.blocked_until:
            remaining_time = st.session_state.blocked_until - datetime.now()
            st.error(f"由于多次验证失败，访问已被临时限制。请在 {remaining_time.seconds//60} 分钟后重试。")
            return False
        else:
            # 解除封禁
            st.session_state.blocked_until = None
            st.session_state.auth_attempts = 0
    
    # 显示验证界面
    st.markdown("## 🔐 访问验证")
    st.markdown("请输入有效的访问密钥以继续使用。")
    
    key_input = st.text_input("访问密钥:", type="password")
    
    if st.button("验证"):
        # 验证密钥
        try:
            key_manager = KeyManager(
                st.secrets["SUPABASE_URL"],
                st.secrets["SUPABASE_KEY"]
            )
            result = key_manager.verify_key(key_input)
            
            if result["valid"]:
                st.session_state.is_authenticated = True
                st.session_state.auth_attempts = 0
                st.success("验证成功！正在进入应用...")
                time.sleep(1)
                st.rerun()
                return True
            else:
                st.session_state.auth_attempts += 1
                st.session_state.last_attempt_time = datetime.now()
                
                remaining_attempts = 3 - st.session_state.auth_attempts
                if remaining_attempts > 0:
                    st.error(f"验证失败: {result['message']}。还剩 {remaining_attempts} 次尝试机会。")
                else:
                    # 设置30分钟封禁期
                    st.session_state.blocked_until = datetime.now() + timedelta(minutes=30)
                    st.error("验证失败次数过多，请30分钟后重试。")
                return False
                
        except Exception as e:
            st.error(f"验证过程出错: {str(e)}")
            return False
    
    return False

def main():
    # 在显示主界面之前进行验证
    if not verify_access():
        return
    
    langsmith_api_key = st.secrets["LANGCHAIN_API_KEY"]
    os.environ["LANGCHAIN_TRACING_V2"] = "true"
    os.environ["LANGCHAIN_API_KEY"] = langsmith_api_key
    os.environ["LANGCHAIN_PROJECT"] = "初稿脑暴平台"
    st.set_page_config(page_title="初稿脑暴助理平台", layout="wide")
    add_custom_css()
    st.markdown("<h1 class='page-title'>初稿脑暴助理</h1>", unsafe_allow_html=True)
    
    if 'prompt_templates' not in st.session_state:
        st.session_state.prompt_templates = PromptTemplates()
    
    tab1, tab2 = st.tabs(["初稿脑暴助理", "提示词设置"])
    st.markdown(f"<div class='model-info'>🤖 当前使用模型: <b>{st.secrets['OPENROUTER_MODEL']}</b></div>", unsafe_allow_html=True)
    
    # 初始化会话状态变量
    if 'document_content' not in st.session_state:
        st.session_state.document_content = None
    if 'transcript_file' not in st.session_state:
        st.session_state.transcript_file = None
    if 'transcript_analysis_done' not in st.session_state:
        st.session_state.transcript_analysis_done = False
    if 'transcript_analysis_result' not in st.session_state:
        st.session_state.transcript_analysis_result = None
    if 'strategist_analysis_done' not in st.session_state:
        st.session_state.strategist_analysis_done = False
    if 'strategist_analysis_result' not in st.session_state:
        st.session_state.strategist_analysis_result = None
    if 'creator_analysis_done' not in st.session_state:
        st.session_state.creator_analysis_done = False
    if 'creator_analysis_result' not in st.session_state:
        st.session_state.creator_analysis_result = None
    if 'show_transcript_analysis' not in st.session_state:
        st.session_state.show_transcript_analysis = False
    if 'show_strategist_analysis' not in st.session_state:
        st.session_state.show_strategist_analysis = False
    if 'show_creator_analysis' not in st.session_state:
        st.session_state.show_creator_analysis = False
    
    with tab1:
        # 添加成绩单上传功能
        transcript_file = st.file_uploader("上传成绩单（可选）", type=['pdf'])
        
        # 添加文件上传功能
        uploaded_file = st.file_uploader("上传初稿文档", type=['docx'])
        
        # 添加选校方案输入框
        school_plan = st.text_area(
            "选校方案",
            value="暂未选校",
            height=100,
            help="请输入已确定的选校方案，包括学校和专业信息"
        )
        
        # 添加自定义需求输入框
        custom_requirements = st.text_area(
            "定制需求（可选）",
            value="无定制需求",
            height=100,
            help="请输入特殊的定制需求，如果没有可以保持默认值"
        )
        
        # 处理上传的成绩单
        if transcript_file is not None:
            st.session_state.transcript_file = transcript_file.read()
            st.success("成绩单上传成功！")
            
            # 添加处理成绩单按钮
            if st.button("处理成绩单", key="process_transcript"):
                st.session_state.show_transcript_analysis = True
                st.session_state.transcript_analysis_done = False
                st.rerun()
        
        # 处理上传的文件
        if uploaded_file is not None:
            document_content = read_docx(uploaded_file.read())
            if document_content:
                st.session_state.document_content = document_content
                st.success("个人陈述上传成功！")
                with st.expander("查看个人陈述内容", expanded=False):
                    st.write(document_content)
            else:
                st.error("无法读取文档内容，请检查文件格式是否正确。")
        
        # 按钮区域 - 始终在顶部
        button_col1, button_col2 = st.columns(2)
        
        with button_col1:
            # 开始背景分析按钮
            if st.button("开始背景分析", key="start_analysis", use_container_width=True):
                if st.session_state.document_content:
                    st.session_state.show_strategist_analysis = True
                    st.session_state.strategist_analysis_done = False
                    st.session_state.creator_analysis_done = False
                    st.session_state.show_creator_analysis = False
                    st.rerun()  # 强制重新运行应用
                else:
                    st.warning("请先上传初稿文档")
        
        with button_col2:
            # 继续内容规划按钮 - 只有在背景分析完成后才启用
            continue_button = st.button(
                "继续内容规划", 
                key="continue_to_creator", 
                disabled=not st.session_state.strategist_analysis_done,
                use_container_width=True
            )
            
            if continue_button:
                st.session_state.show_creator_analysis = True
                st.session_state.creator_analysis_done = False
                st.rerun()  # 强制重新运行应用
        
        # 创建结果显示区域
        results_container = st.container()
        
        # 显示成绩单分析（如果需要）
        if st.session_state.show_transcript_analysis:
            with results_container:
                st.markdown("---")
                st.subheader("📊 成绩单分析")
                
                if not st.session_state.transcript_analysis_done:
                    try:
                        transcript_analyzer = TranscriptAnalyzer(
                            api_key=st.secrets["OPENROUTER_API_KEY"],  # 使用OpenRouter API密钥
                            prompt_templates=st.session_state.prompt_templates
                        )
                        
                        with st.spinner("正在分析成绩单..."):
                            # 处理成绩单分析
                            result = transcript_analyzer.analyze_transcript(
                                st.session_state.transcript_file, 
                                school_plan
                            )
                            
                            if result["status"] == "success":
                                # 保存成绩单分析结果到 session_state
                                st.session_state.transcript_analysis_result = result["transcript_analysis"]
                                st.session_state.transcript_analysis_done = True
                                st.success("✅ 成绩单分析完成！")
                            else:
                                st.error(f"成绩单分析出错: {result['message']}")
                    
                    except Exception as e:
                        st.error(f"处理过程中出错: {str(e)}")
                else:
                    # 如果已经完成，直接显示结果
                    st.markdown(st.session_state.transcript_analysis_result)
                    st.success("✅ 成绩单分析完成！")
        
        # 显示背景分析（如果需要）
        if st.session_state.show_strategist_analysis:
            with results_container:
                st.markdown("---")
                st.subheader("📊 第一阶段：背景分析")
                
                if not st.session_state.strategist_analysis_done:
                    try:
                        agent = BrainstormingAgent(
                            api_key=st.secrets["OPENROUTER_API_KEY"],
                            prompt_templates=st.session_state.prompt_templates
                        )
                        
                        with st.spinner("正在进行背景分析..."):
                            # 获取成绩单分析结果（如果有）
                            transcript_analysis = ""
                            if st.session_state.transcript_analysis_done:
                                transcript_analysis = st.session_state.transcript_analysis_result
                            
                            # 处理第一阶段分析
                            result = agent.process_strategist(
                                st.session_state.document_content, 
                                school_plan,
                                transcript_analysis
                            )
                            
                            if result["status"] == "success":
                                # 保存策略分析结果到 session_state
                                st.session_state.strategist_analysis_result = result["strategist_analysis"]
                                st.session_state.strategist_analysis_done = True
                                st.success("✅ 背景分析完成！")
                                st.rerun()  # 强制重新运行应用以更新按钮状态
                            else:
                                st.error(f"背景分析出错: {result['message']}")
                    
                    except Exception as e:
                        st.error(f"处理过程中出错: {str(e)}")
                else:
                    # 如果已经完成，直接显示结果
                    st.markdown(st.session_state.strategist_analysis_result)
                    st.success("✅ 背景分析完成！")
        
        # 显示内容规划（如果需要）
        if st.session_state.show_creator_analysis:
            with results_container:
                st.markdown("---")
                st.subheader("📝 第二阶段：内容规划")
                
                if not st.session_state.creator_analysis_done:
                    try:
                        agent = BrainstormingAgent(
                            api_key=st.secrets["OPENROUTER_API_KEY"],
                            prompt_templates=st.session_state.prompt_templates
                        )
                        
                        with st.spinner("正在进行内容规划..."):
                            creator_result = agent.process_creator(
                                st.session_state.strategist_analysis_result,
                                school_plan,
                                st.session_state.transcript_analysis_result,
                                custom_requirements
                            )
                            
                            if creator_result["status"] == "success":
                                st.session_state.creator_analysis_result = creator_result["creator_output"]
                                st.session_state.creator_analysis_done = True
                                st.success("✅ 内容规划完成！")
                            else:
                                st.error(f"内容规划出错: {creator_result['message']}")
                    
                    except Exception as e:
                        st.error(f"处理过程中出错: {str(e)}")
                else:
                    # 如果已经完成，直接显示结果
                    st.markdown(st.session_state.creator_analysis_result)
                    st.success("✅ 内容规划完成！")
    
    with tab2:
        st.title("提示词设置")
        
        prompt_templates = st.session_state.prompt_templates
        
        # 成绩单分析设置
        st.subheader("成绩单分析")
        transcript_role = st.text_area(
            "角色设定",
            value=prompt_templates.get_template('transcript_role'),
            height=200,
            key="transcript_role"
        )
        
        transcript_task = st.text_area(
            "任务说明",
            value=prompt_templates.get_template('transcript_task'),
            height=200,
            key="transcript_task"
        )
        
        transcript_output = st.text_area(
            "输出格式",
            value=prompt_templates.get_template('transcript_output'),
            height=200,
            key="transcript_output"
        )
        
        # Agent 1 设置
        st.subheader("Agent 1 - 档案策略师")
        consultant_role1 = st.text_area(
            "角色设定",
            value=prompt_templates.get_template('consultant_role1'),
            height=200,
            key="consultant_role1"
        )
        
        consultant_task1 = st.text_area(
            "任务说明",
            value=prompt_templates.get_template('consultant_task1'),
            height=200,
            key="consultant_task1"
        )

        output_format1 = st.text_area(
            "输出格式",
            value=prompt_templates.get_template('output_format1'),
            height=200,
            key="output_format1"
        )
        # Agent 2 设置
        st.subheader("Agent 2 - 内容创作师")
        consultant_role2 = st.text_area(
            "角色设定",
            value=prompt_templates.get_template('consultant_role2'),
            height=200,
            key="consultant_role2"
        )

        consultant_task2 = st.text_area(
            "任务说明",
            value=prompt_templates.get_template('consultant_task2'),
            height=200,
            key="consultant_task2"
        )

        output_format2 = st.text_area(
            "输出格式",
            value=prompt_templates.get_template('output_format2'),
            height=200,
            key="output_format2"
        )
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("更新提示词", key="update_prompts"):
                prompt_templates.update_template('transcript_role', transcript_role)
                prompt_templates.update_template('transcript_task', transcript_task)
                prompt_templates.update_template('transcript_output', transcript_output)
                prompt_templates.update_template('consultant_role1', consultant_role1)
                prompt_templates.update_template('output_format1', output_format1)
                prompt_templates.update_template('consultant_task1', consultant_task1)
                prompt_templates.update_template('consultant_role2', consultant_role2)
                prompt_templates.update_template('output_format2', output_format2)
                prompt_templates.update_template('consultant_task2', consultant_task2)
                st.success("✅ 提示词已更新！")
        
        with col2:
            if st.button("重置为默认提示词", key="reset_prompts"):
                prompt_templates.reset_to_default()
                st.rerun()

if __name__ == "__main__":
    main()
